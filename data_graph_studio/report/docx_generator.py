"""
Word Document (DOCX) Report Generator
Word 문서 레포트 생성기

Uses python-docx for document generation.
Supports styles, tables, images, and table of contents.
"""

import io
import base64
import logging

from data_graph_studio.core.report import (
    ReportGenerator,
    ReportData,
    ReportOptions,
    PageSize,
    PageOrientation,
)

logger = logging.getLogger(__name__)

# Check if python-docx is available
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DOCX_AVAILABLE = True
except ImportError:
    pass


class DOCXReportGenerator(ReportGenerator):
    """Word 문서 레포트 생성기"""

    # Page sizes in cm
    PAGE_SIZES = {
        PageSize.A4: (21.0, 29.7),
        PageSize.LETTER: (21.59, 27.94),
        PageSize.LEGAL: (21.59, 35.56),
        PageSize.A3: (29.7, 42.0),
    }

    def generate(self, report_data: ReportData, options: ReportOptions) -> bytes:
        """Word 문서 생성"""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not available. Install it with:\n"
                "  pip install python-docx"
            )

        doc = Document()

        # 문서 설정
        self._setup_document(doc, options)

        # 스타일 설정
        self._setup_styles(doc)

        # 헤더 추가
        self._add_header(doc, report_data, options)

        # 목차 추가
        self._add_table_of_contents(doc, options)

        # Executive Summary
        if options.include_executive_summary:
            self._add_executive_summary(doc, report_data, options)

        # Data Overview
        if options.include_data_overview:
            self._add_data_overview(doc, report_data, options)

        # Statistics
        if options.include_statistics:
            self._add_statistics(doc, report_data, options)

        # Visualizations
        if options.include_visualizations and report_data.charts:
            self._add_visualizations(doc, report_data, options)

        # Comparison Analysis
        if options.include_comparison and report_data.is_multi_dataset():
            self._add_comparison(doc, report_data, options)

        # Data Tables
        if options.include_tables and report_data.tables:
            self._add_tables(doc, report_data, options)

        # Appendix
        if options.include_appendix:
            self._add_appendix(doc, report_data, options)

        # 바이트로 저장
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _hex_to_rgb(self, hex_color: str) -> "RGBColor":
        """HEX를 RGBColor로 변환"""
        hex_color = hex_color.lstrip("#")
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)

    def _setup_document(self, doc: "Document", options: ReportOptions):
        """문서 설정"""
        section = doc.sections[0]

        # 페이지 크기
        page_size = self.PAGE_SIZES.get(options.page_size, (21.0, 29.7))

        if options.orientation == PageOrientation.LANDSCAPE:
            section.page_width = Cm(page_size[1])
            section.page_height = Cm(page_size[0])
        else:
            section.page_width = Cm(page_size[0])
            section.page_height = Cm(page_size[1])

        # 마진
        margins = options.margins
        section.top_margin = Cm(margins.get("top", 2.54))
        section.bottom_margin = Cm(margins.get("bottom", 2.54))
        section.left_margin = Cm(margins.get("left", 2.54))
        section.right_margin = Cm(margins.get("right", 2.54))

        # 헤더/푸터
        if options.header_text:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = options.header_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if options.footer_text:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = options.footer_text
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _setup_styles(self, doc: "Document"):
        """스타일 설정"""
        styles = doc.styles

        # 제목 스타일
        if "Report Title" not in [s.name for s in styles]:
            title_style = styles.add_style("Report Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = self._hex_to_rgb(self.template.primary_color)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)

        # 섹션 제목 스타일
        if "Section Title" not in [s.name for s in styles]:
            section_style = styles.add_style("Section Title", WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.font.color.rgb = self._hex_to_rgb(self.template.primary_color)
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(12)

        # 서브섹션 스타일
        if "Subsection Title" not in [s.name for s in styles]:
            subsection_style = styles.add_style(
                "Subsection Title", WD_STYLE_TYPE.PARAGRAPH
            )
            subsection_style.font.size = Pt(13)
            subsection_style.font.bold = True
            subsection_style.paragraph_format.space_before = Pt(12)
            subsection_style.paragraph_format.space_after = Pt(6)

    def _add_header(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """문서 헤더 (제목 페이지)"""
        meta = report_data.metadata

        # 로고 (있는 경우)
        if meta.logo_base64:
            try:
                logo_bytes = base64.b64decode(meta.logo_base64)
                logo_stream = io.BytesIO(logo_bytes)
                doc.add_picture(logo_stream, width=Inches(2))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Failed to add logo: {e}")

        # 제목
        title_para = doc.add_paragraph(meta.title)
        title_para.style = "Report Title"

        # 부제목
        if meta.subtitle:
            subtitle = doc.add_paragraph(meta.subtitle)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(14)
            subtitle.runs[0].font.italic = True

        # 메타 정보
        doc.add_paragraph()  # 빈 줄

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_text = f"Generated: {meta.created_at.strftime('%Y-%m-%d %H:%M')}"
        if meta.author:
            meta_text += f"\nAuthor: {meta.author}"
        meta_text += f"\nData Graph Studio v{meta.version}"

        meta_para.add_run(meta_text).font.size = Pt(10)

        # 페이지 나누기
        doc.add_page_break()

    def _add_table_of_contents(self, doc: "Document", options: ReportOptions):
        """목차 추가"""
        is_ko = options.language == "ko"
        toc_title = "목차" if is_ko else "Table of Contents"

        doc.add_paragraph(toc_title, style="Section Title")

        # 목차 필드 (Word에서 자동 업데이트)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # TOC 필드 코드
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)

        # 플레이스홀더 텍스트
        placeholder = doc.add_paragraph(
            "(Update this table of contents in Word: Right-click → Update Field)"
        )
        placeholder.runs[0].font.italic = True
        placeholder.runs[0].font.size = Pt(9)

        run._r.append(fld_char_end)

        doc.add_page_break()

    def _add_section_heading(self, doc: "Document", title: str, level: int = 1):
        """섹션 제목 추가"""
        if level == 1:
            heading = doc.add_heading(title, level=1)
            heading.style = doc.styles["Heading 1"]
        elif level == 2:
            heading = doc.add_heading(title, level=2)
        else:
            doc.add_paragraph(title, style="Subsection Title")

    def _add_executive_summary(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """Executive Summary 추가"""
        is_ko = options.language == "ko"
        title = "요약" if is_ko else "Executive Summary"

        self._add_section_heading(doc, title, 1)

        # 주요 지표 테이블
        metrics = [
            (
                "Total Rows" if not is_ko else "총 행 수",
                self.format_number(report_data.get_total_rows(), 0),
            ),
            ("Datasets" if not is_ko else "데이터셋", str(len(report_data.datasets))),
        ]

        if report_data.datasets:
            metrics.append(
                (
                    "Columns" if not is_ko else "컬럼 수",
                    str(report_data.datasets[0].column_count),
                )
            )

        if report_data.charts:
            metrics.append(
                ("Charts" if not is_ko else "차트", str(len(report_data.charts)))
            )

        # 2열 테이블로 메트릭 표시
        table = doc.add_table(rows=1, cols=len(metrics))
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for i, (label, value) in enumerate(metrics):
            cell = header_cells[i]
            cell.text = f"{label}\n{value}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 스타일 적용
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
                para.runs[0].font.bold = True if "\n" not in para.text else False

        doc.add_paragraph()  # 공백

        # Key Findings
        if report_data.key_findings:
            findings_title = "핵심 발견 사항" if is_ko else "Key Findings"
            self._add_section_heading(doc, findings_title, 2)

            for finding in report_data.key_findings:
                para = doc.add_paragraph(finding, style="List Bullet")

        # Recommendations
        if report_data.recommendations:
            rec_title = "권장 사항" if is_ko else "Recommendations"
            self._add_section_heading(doc, rec_title, 2)

            for rec in report_data.recommendations:
                para = doc.add_paragraph(rec, style="List Bullet")

    def _add_data_overview(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """Data Overview 추가"""
        is_ko = options.language == "ko"
        title = "데이터 개요" if is_ko else "Data Overview"

        self._add_section_heading(doc, title, 1)

        for ds in report_data.datasets:
            self._add_section_heading(doc, ds.name, 2)

            # 데이터셋 정보 테이블
            info_items = [
                ("Rows" if not is_ko else "행 수", self.format_number(ds.row_count, 0)),
                ("Columns" if not is_ko else "컬럼 수", str(ds.column_count)),
                (
                    "Memory" if not is_ko else "메모리",
                    self.format_bytes(ds.memory_bytes),
                ),
            ]

            if ds.file_path:
                info_items.insert(
                    0, ("File" if not is_ko else "파일", ds.file_path.split("/")[-1])
                )

            if ds.date_range:
                date_str = f"{ds.date_range['min']} ~ {ds.date_range['max']}"
                info_items.append(
                    ("Date Range" if not is_ko else "날짜 범위", date_str)
                )

            table = doc.add_table(rows=len(info_items), cols=2)
            table.style = "Table Grid"

            for i, (label, value) in enumerate(info_items):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[1].text = value

            doc.add_paragraph()  # 공백

            # 컬럼 목록
            if ds.columns:
                columns_title = "컬럼 목록" if is_ko else "Columns"
                doc.add_paragraph(columns_title, style="Subsection Title")

                # 컬럼 테이블
                col_table = doc.add_table(rows=1, cols=3)
                col_table.style = "Table Grid"

                headers = ["Column", "Type", "Missing"]
                for i, header in enumerate(headers):
                    col_table.rows[0].cells[i].text = header
                    col_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

                for col in ds.columns[:20]:  # 최대 20개
                    row = col_table.add_row()
                    row.cells[0].text = col
                    row.cells[1].text = ds.column_types.get(col, "unknown")
                    row.cells[2].text = str(ds.missing_values.get(col, 0))

                if len(ds.columns) > 20:
                    doc.add_paragraph(f"... and {len(ds.columns) - 20} more columns")

                doc.add_paragraph()

    def _add_statistics(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """Statistics 추가"""
        is_ko = options.language == "ko"
        title = "통계 요약" if is_ko else "Statistical Summary"

        self._add_section_heading(doc, title, 1)

        for dataset_id, stats_list in report_data.statistics.items():
            if not stats_list:
                continue

            # 데이터셋 이름 찾기
            dataset_name = dataset_id
            for ds in report_data.datasets:
                if ds.id == dataset_id:
                    dataset_name = ds.name
                    break

            self._add_section_heading(doc, dataset_name, 2)

            # 통계 테이블
            headers = [
                "Column",
                "Count",
                "Mean",
                "Median",
                "Std",
                "Min",
                "Max",
                "Q1",
                "Q3",
            ]

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"

            # 헤더
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)

            # 데이터
            for stat in stats_list[:30]:  # 최대 30행
                row = table.add_row()
                row.cells[0].text = stat.column[:20]
                row.cells[1].text = str(stat.count)
                row.cells[2].text = self.format_number(stat.mean, 2)
                row.cells[3].text = self.format_number(stat.median, 2)
                row.cells[4].text = self.format_number(stat.std, 2)
                row.cells[5].text = self.format_number(stat.min, 2)
                row.cells[6].text = self.format_number(stat.max, 2)
                row.cells[7].text = self.format_number(stat.q1, 2)
                row.cells[8].text = self.format_number(stat.q3, 2)

                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)

            doc.add_paragraph()

    def _add_visualizations(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """Visualizations 추가"""
        is_ko = options.language == "ko"
        title = "시각화" if is_ko else "Visualizations"

        self._add_section_heading(doc, title, 1)

        for chart in report_data.charts:
            self._add_section_heading(doc, chart.title, 2)

            # 이미지 삽입
            if chart.image_bytes:
                try:
                    image_stream = io.BytesIO(chart.image_bytes)
                    doc.add_picture(image_stream, width=Inches(5.5))

                    # 이미지 중앙 정렬
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                except Exception as e:
                    logger.warning(f"Failed to add chart image: {e}")
                    doc.add_paragraph("[Chart image not available]")
            elif chart.image_base64:
                try:
                    image_bytes = base64.b64decode(chart.image_base64)
                    image_stream = io.BytesIO(image_bytes)
                    doc.add_picture(image_stream, width=Inches(5.5))

                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                except Exception as e:
                    logger.warning(f"Failed to add chart image: {e}")
                    doc.add_paragraph("[Chart image not available]")

            # 설명
            if chart.description:
                desc_para = doc.add_paragraph(chart.description)
                desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                desc_para.runs[0].font.italic = True
                desc_para.runs[0].font.size = Pt(10)

            doc.add_paragraph()

    def _add_comparison(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """Comparison Analysis 추가"""
        is_ko = options.language == "ko"
        title = "비교 분석" if is_ko else "Comparison Analysis"

        self._add_section_heading(doc, title, 1)

        # Statistical Test Results
        if report_data.comparisons:
            test_title = "통계 검정 결과" if is_ko else "Statistical Test Results"
            self._add_section_heading(doc, test_title, 2)

            for comp in report_data.comparisons:
                # 검정 제목
                doc.add_paragraph(
                    f"{comp.test_type}: {comp.dataset_a_name} vs {comp.dataset_b_name}",
                    style="Subsection Title",
                )

                # 결과 테이블
                results = [
                    ("Column", comp.column),
                    ("Test Statistic", f"{comp.test_statistic:.4f}"),
                    ("p-value", f"{comp.p_value:.4f}"),
                ]

                if comp.effect_size is not None:
                    results.append(
                        (
                            "Effect Size",
                            f"{comp.effect_size:.3f} ({comp.effect_size_interpretation})",
                        )
                    )

                # 유의성
                sig_text = "Not Significant"
                if comp.p_value < 0.001:
                    sig_text = "*** Highly Significant (p < 0.001)"
                elif comp.p_value < 0.01:
                    sig_text = "** Very Significant (p < 0.01)"
                elif comp.p_value < 0.05:
                    sig_text = "* Significant (p < 0.05)"

                results.append(("Significance", sig_text))

                table = doc.add_table(rows=len(results), cols=2)
                table.style = "Table Grid"

                for i, (label, value) in enumerate(results):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                    table.rows[i].cells[1].text = value

                # 해석
                if comp.interpretation:
                    interp_para = doc.add_paragraph(comp.interpretation)
                    interp_para.runs[0].font.italic = True

                doc.add_paragraph()

        # Difference Analysis
        if report_data.differences:
            diff_title = "차이 분석" if is_ko else "Difference Analysis"
            self._add_section_heading(doc, diff_title, 2)

            for diff in report_data.differences:
                doc.add_paragraph(
                    f"{diff.dataset_a_name} vs {diff.dataset_b_name}",
                    style="Subsection Title",
                )

                # 차이 요약 테이블
                summary = [
                    ("Matched Records", str(diff.matched_records)),
                    (
                        "Positive (A > B)",
                        f"{diff.positive_count} ({diff.positive_percentage:.1f}%)",
                    ),
                    (
                        "Negative (A < B)",
                        f"{diff.negative_count} ({diff.negative_percentage:.1f}%)",
                    ),
                    (
                        "No Change",
                        f"{diff.neutral_count} ({diff.neutral_percentage:.1f}%)",
                    ),
                    ("Total Difference", self.format_number(diff.total_difference)),
                    ("Mean Difference", self.format_number(diff.mean_difference)),
                ]

                table = doc.add_table(rows=len(summary), cols=2)
                table.style = "Table Grid"

                for i, (label, value) in enumerate(summary):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                    table.rows[i].cells[1].text = value

                doc.add_paragraph()

    def _add_tables(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """Data Tables 추가"""
        is_ko = options.language == "ko"
        title = "데이터 테이블" if is_ko else "Data Tables"

        self._add_section_heading(doc, title, 1)

        for table_data in report_data.tables:
            self._add_section_heading(doc, table_data.title, 2)

            # 테이블 생성
            num_cols = min(len(table_data.columns), 8)  # 최대 8열
            num_rows = min(len(table_data.rows), 25) + 1  # 헤더 + 최대 25행

            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = "Table Grid"

            # 헤더
            for i, col in enumerate(table_data.columns[:num_cols]):
                cell = table.rows[0].cells[i]
                cell.text = col[:15]  # 최대 15자
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)

            # 데이터
            for row_idx, row_data in enumerate(table_data.rows[:25]):
                for col_idx, cell_value in enumerate(row_data[:num_cols]):
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell_text = str(cell_value) if cell_value is not None else "-"
                    if len(cell_text) > 20:
                        cell_text = cell_text[:17] + "..."
                    cell.text = cell_text
                    cell.paragraphs[0].runs[0].font.size = Pt(8)

            # 페이지네이션 정보
            if table_data.total_rows > 25:
                doc.add_paragraph(f"Showing 25 of {table_data.total_rows} rows").runs[
                    0
                ].font.italic = True

            doc.add_paragraph()

    def _add_appendix(
        self, doc: "Document", report_data: ReportData, options: ReportOptions
    ):
        """Appendix 추가"""
        is_ko = options.language == "ko"
        title = "부록" if is_ko else "Appendix"

        self._add_section_heading(doc, title, 1)

        # Methodology
        if report_data.methodology_notes:
            method_title = "분석 방법론" if is_ko else "Methodology"
            self._add_section_heading(doc, method_title, 2)

            for note in report_data.methodology_notes:
                doc.add_paragraph(note, style="List Bullet")

        # Data Quality Notes
        if report_data.data_quality_notes:
            quality_title = "데이터 품질 노트" if is_ko else "Data Quality Notes"
            self._add_section_heading(doc, quality_title, 2)

            for note in report_data.data_quality_notes:
                doc.add_paragraph(note, style="List Bullet")
